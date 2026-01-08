"""
Script untuk membuat dokumen Tahap Testing (Design Thinking) dalam format DOCX
Untuk UAS Mata Kuliah Design Thinking
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

# Create document
doc = Document()

# Set document title
title = doc.add_heading('D. Tahap Testing (Test Phase)', level=1)
title.alignment = WD_ALIGN_PARAGRAPH.LEFT

# Context paragraph
doc.add_paragraph(
    'Tahap Testing merupakan fase kelima dan terakhir dalam proses Design Thinking. '
    'Pada tahap ini, prototype yang telah dibuat diuji kepada pengguna untuk mendapatkan '
    'feedback dan validasi terhadap solusi yang diusulkan. Hasil pengujian akan digunakan '
    'untuk iterasi dan penyempurnaan produk.'
)

doc.add_paragraph()

# ============ SECTION 1: Tujuan Pengujian ============
doc.add_heading('1. Tujuan Pengujian', level=2)

doc.add_paragraph(
    'Pengujian pada prototype Notes App (Aplikasi Manajemen Tugas Desktop) dilakukan '
    'dengan beberapa tujuan sebagai berikut:'
)

# Create table for objectives
table1 = doc.add_table(rows=6, cols=3)
table1.style = 'Table Grid'

# Header row
header_cells = table1.rows[0].cells
header_cells[0].text = 'No'
header_cells[1].text = 'Tujuan Pengujian'
header_cells[2].text = 'Deskripsi'

# Data rows
data = [
    ('1', 'Menguji Alur Solusi (Solution Flow)', 
     'Memvalidasi apakah alur aplikasi sudah menjawab pain points pengguna: dari membuka app → menambah tugas → mendapat reminder → menyelesaikan tugas'),
    ('2', 'Menguji Kemudahan Penggunaan (Usability)', 
     'Mengukur seberapa mudah pengguna (mahasiswa) melakukan task utama tanpa bantuan, seperti menambah tugas, mengedit, dan menandai selesai'),
    ('3', 'Menguji Pengalaman Pengguna (User Experience)', 
     'Mengevaluasi kenyamanan visual, responsivitas interface, dan kepuasan pengguna terhadap Kanban board, Calendar view, dan AI Chat assistant'),
    ('4', 'Menguji Pemahaman Pengguna terhadap Solusi', 
     'Memastikan pengguna memahami value proposition: sinkronisasi cloud, integrasi Telegram Bot, dan fitur "Ask with Kanee" AI untuk bantuan akademik'),
    ('5', 'Mengumpulkan Feedback untuk Iterasi', 
     'Mendapatkan insight dan saran perbaikan dari pengguna nyata untuk menyempurnakan prototype sebelum implementasi final'),
]

for i, row_data in enumerate(data):
    row = table1.rows[i + 1].cells
    row[0].text = row_data[0]
    row[1].text = row_data[1]
    row[2].text = row_data[2]

doc.add_paragraph()  # Spacer

# ============ SECTION 2: Metode Testing ============
doc.add_heading('2. Metode Testing yang Digunakan', level=2)

doc.add_paragraph(
    'Metode testing yang dipilih disesuaikan dengan prinsip Design Thinking yang berfokus pada '
    'human-centered design. Kombinasi metode berikut digunakan untuk mendapatkan feedback yang komprehensif:'
)

# Method A: Usability Testing
doc.add_heading('A. Usability Testing dengan Real Users', level=3)

p1 = doc.add_paragraph()
p1.add_run('Deskripsi: ').bold = True
p1.add_run('Mengundang 5 mahasiswa sebagai pengguna target untuk mencoba prototype secara langsung. '
           'Metode ini sesuai dengan prinsip "Empathy" dalam Design Thinking karena melibatkan pengguna nyata.')

doc.add_paragraph()
p2 = doc.add_paragraph()
p2.add_run('Skenario Pengujian:').bold = True

scenarios = doc.add_paragraph(style='List Bullet')
scenarios.add_run('Skenario 1: Tambah tugas baru dengan deadline "besok jam 10 malam"')

doc.add_paragraph('Skenario 2: Edit deskripsi tugas yang sudah ada', style='List Bullet')
doc.add_paragraph('Skenario 3: Tandai tugas sebagai selesai (completed)', style='List Bullet')
doc.add_paragraph('Skenario 4: Sinkronisasi data dengan cloud (Sync)', style='List Bullet')
doc.add_paragraph('Skenario 5: Tanya AI helper tentang rumus matematika', style='List Bullet')

doc.add_paragraph()
p3 = doc.add_paragraph()
p3.add_run('Metrik Pengukuran:').bold = True

doc.add_paragraph('Task Completion Rate: Persentase tugas yang berhasil diselesaikan', style='List Bullet')
doc.add_paragraph('Time on Task: Waktu yang dibutuhkan untuk menyelesaikan setiap skenario', style='List Bullet')
doc.add_paragraph('Error Rate: Jumlah kesalahan yang dilakukan pengguna', style='List Bullet')
doc.add_paragraph('Satisfaction Score: Tingkat kepuasan pengguna (skala 1-5)', style='List Bullet')

# Method B: Think-Aloud Protocol
doc.add_heading('B. Think-Aloud Protocol', level=3)

p4 = doc.add_paragraph()
p4.add_run('Deskripsi: ').bold = True
p4.add_run('Pengguna diminta untuk menyuarakan pikiran, perasaan, dan kebingungan mereka secara real-time '
           'saat berinteraksi dengan prototype. Metode ini memberikan insight mendalam tentang mental model pengguna.')

doc.add_paragraph()
p5 = doc.add_paragraph()
p5.add_run('Tujuan: ').bold = True
p5.add_run('Mengidentifikasi:')

doc.add_paragraph('Bagian interface yang membingungkan', style='List Bullet')
doc.add_paragraph('Fitur yang tidak intuitif', style='List Bullet')
doc.add_paragraph('Ekspektasi pengguna yang tidak terpenuhi', style='List Bullet')
doc.add_paragraph('Momen "aha!" atau kepuasan pengguna', style='List Bullet')

# Method C: Observation
doc.add_heading('C. Observation (Pengamatan Langsung)', level=3)

p6 = doc.add_paragraph()
p6.add_run('Deskripsi: ').bold = True
p6.add_run('Tim melakukan observasi langsung terhadap perilaku pengguna saat menggunakan prototype, '
           'mencatat gesture, ekspresi wajah, dan pola interaksi.')

doc.add_paragraph()
p7 = doc.add_paragraph()
p7.add_run('Hal yang Diamati: ').bold = True

doc.add_paragraph('Apakah pengguna ragu-ragu sebelum klik?', style='List Bullet')
doc.add_paragraph('Apakah ada gesture frustasi (mengeluh, mengulang aksi)?', style='List Bullet')
doc.add_paragraph('Bagaimana reaksi saat fitur berhasil/gagal?', style='List Bullet')

# Method D: SUS Questionnaire
doc.add_heading('D. System Usability Scale (SUS) Questionnaire', level=3)

p8 = doc.add_paragraph()
p8.add_run('Deskripsi: ').bold = True
p8.add_run('Kuesioner standar industri yang terdiri dari 10 pertanyaan untuk mengukur persepsi usability '
           'secara kuantitatif.')

doc.add_paragraph()
p9 = doc.add_paragraph()
p9.add_run('Skala Penilaian: ').bold = True
p9.add_run('1 (Sangat Tidak Setuju) - 5 (Sangat Setuju)')

doc.add_paragraph()
p10 = doc.add_paragraph()
p10.add_run('Target Score: ').bold = True
p10.add_run('> 68 (di atas rata-rata industri)')

doc.add_paragraph()

# ============ SECTION 3: Hasil dan Iterasi ============
doc.add_heading('3. Rencana Iterasi Berdasarkan Hasil Testing', level=2)

doc.add_paragraph(
    'Sesuai dengan prinsip Design Thinking yang iteratif, hasil testing akan digunakan untuk:'
)

doc.add_paragraph('Mengidentifikasi area perbaikan pada interface dan user flow', style='List Bullet')
doc.add_paragraph('Memprioritaskan fitur berdasarkan feedback pengguna', style='List Bullet')
doc.add_paragraph('Melakukan iterasi desain sebelum development lanjutan', style='List Bullet')
doc.add_paragraph('Validasi ulang dengan pengguna setelah perbaikan dilakukan', style='List Bullet')

# Save document
output_path = r'e:\Personal project\Notes-app-desktop\UAS_Design_Thinking_Tahap_Testing.docx'
doc.save(output_path)

print(f"Dokumen UAS berhasil dibuat: {output_path}")
